"""Document parser for various file formats."""

import os
from pathlib import Path
from typing import Optional, Dict, Any, List
import aiofiles
from pypdf import PdfReader
from docx import Document
from PIL import Image
import io
import cv2
import tempfile
import zipfile
import tarfile
import shutil
from typing import Optional, Dict, Any, List, Union


class DocumentParser:
    """Parser for various document formats."""

    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".docx",
        ".doc",
        ".txt",
        ".md",
        ".log",
        ".png",
        ".jpg",
        ".jpeg",
        ".mp4",
        ".mov",
        ".avi",
        ".zip",
        ".tar",
        ".gz",
    }

    SECURITY_KEYWORDS = {
        "iam",
        "policy",
        "auth",
        "encrypt",
        "secret",
        "access",
        "role",
        "allow",
        "deny",
        "ingress",
        "egress",
        "mfa",
        "tls",
        "ssl",
        "vault",
        "rbac",
        "kms",
        "security_group",
        "firewall",
        "authorization",
        "authentication",
    }

    def __init__(self, upload_dir: str = "./uploads"):
        """Initialize document parser."""
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    async def parse(self, file_path: Path) -> Optional[Union[Dict[str, Any], List[Dict[str, Any]]]]:
        """Parse a document and extract content."""
        if not file_path.exists():
            return None

        extension = file_path.suffix.lower()

        if extension == ".pdf":
            return await self._parse_pdf(file_path)
        elif extension in {".docx", ".doc"}:
            return await self._parse_docx(file_path)
        elif extension in {".txt", ".md"}:
            return await self._parse_text(file_path, doc_type="text")
        elif extension in {".log"}:
            return await self._parse_text(file_path, doc_type="log")
        elif extension in {".png", ".jpg", ".jpeg"}:
            return await self._parse_image(file_path)
        elif extension in {".mp4", ".mov", ".avi"}:
            return await self._parse_video(file_path)
        elif extension in {".zip", ".tar", ".gz"}:
            return await self._parse_archive(file_path)
        else:
            return None

    async def _parse_archive(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extract and parse files from a repository archive."""
        extracted_files = []
        temp_extract_dir = Path(tempfile.mkdtemp(prefix="repo_extract_"))

        try:
            # Extract the archive
            if file_path.suffix.lower() == ".zip":
                with zipfile.ZipFile(file_path, "r") as zip_ref:
                    zip_ref.extractall(temp_extract_dir)
            elif file_path.suffix.lower() in {".tar", ".gz"}:
                mode = (
                    "r:gz"
                    if file_path.suffix.lower() == ".gz" or file_path.name.endswith(".tar.gz")
                    else "r:"
                )
                with tarfile.open(file_path, mode) as tar_ref:
                    tar_ref.extractall(temp_extract_dir)

            # Walk through the extracted files
            ignore_dirs = {
                ".git",
                ".terraform",
                "node_modules",
                "venv",
                "__pycache__",
                "obj",
                "bin",
                "dist",
            }
            ignore_exts = {
                ".exe",
                ".dll",
                ".so",
                ".pyc",
                ".png",
                ".jpg",
                ".jpeg",
                ".gif",
                ".pdf",
                ".docx",
                ".zip",
                ".tar",
                ".gz",
            }

            for root, dirs, files in os.walk(temp_extract_dir):
                # Filter out ignored directories
                dirs[:] = [d for d in dirs if d not in ignore_dirs]

                for name in files:
                    full_path = Path(root) / name

                    # Skip large files > 1MB
                    if full_path.stat().st_size > 1 * 1024 * 1024:
                        continue

                    if full_path.suffix.lower() in ignore_exts:
                        continue

                    # Determine document type (IaC vs Code)
                    doc_type = self._classify_file(full_path)
                    if not doc_type:
                        continue

                    # Read content
                    try:
                        async with aiofiles.open(
                            full_path, "r", encoding="utf-8", errors="ignore"
                        ) as f:
                            content = await f.read()
                    except Exception:
                        continue

                    if not content.strip():
                        continue

                    # Balanced security keyword check
                    contains_keywords = any(kw in content.lower() for kw in self.SECURITY_KEYWORDS)

                    # Store relative path within repo
                    rel_path = full_path.relative_to(temp_extract_dir)

                    extracted_files.append(
                        {
                            "type": doc_type,
                            "filename": str(rel_path),
                            "original_archive": file_path.name,
                            "full_text": content,
                            "contains_security_keywords": contains_keywords,
                            "message": f"Extracted from {file_path.name}",
                        }
                    )

        finally:
            # We keep the temp dir for a moment? No, better copy what we need and delete.
            # But here we return the text content, so we can delete.
            shutil.rmtree(temp_extract_dir, ignore_errors=True)

        # Sort: prioritize files with security keywords
        extracted_files.sort(key=lambda x: x["contains_security_keywords"], reverse=True)

        return extracted_files

    def _classify_file(self, file_path: Path) -> Optional[str]:
        """Classify a file as IaC, Code, or other."""
        ext = file_path.suffix.lower()
        name = file_path.name.lower()

        # Infrastructure as Code (Tier 2)
        iac_exts = {".tf", ".tfvars", ".yaml", ".yml", ".json"}
        iac_names = {"dockerfile", "chart.yaml", "values.yaml", "kustomization.yaml"}

        if ext in iac_exts or name in iac_names:
            # Special check for Helm/K8s/Terraform
            if ext in {".yaml", ".yml", ".json"}:
                # Only include if security-relevant or part of a known pattern
                # (Balanced: include if in a 'templates' dir or has keywords)
                # For now, let's be more inclusive for .tf and known names
                if ext == ".tf" or name in iac_names:
                    return "iac"

                # Check for Helm templates or K8s
                if "templates/" in str(file_path).lower():
                    return "iac"

                return "iac"  # Return iac anyway, we'll filter by keyword in archive loop if needed
            return "iac"

        # Source Code (Tier 4)
        code_exts = {".py", ".js", ".ts", ".go", ".java", ".cs", ".rb", ".php", ".c", ".cpp", ".h"}
        if ext in code_exts:
            return "code"

        return None

    async def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF document."""
        reader = PdfReader(str(file_path))

        pages = []
        full_text = []

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append({"page_number": i + 1, "text": text, "char_count": len(text)})
            full_text.append(text)

        return {
            "type": "pdf",
            "filename": file_path.name,
            "page_count": len(reader.pages),
            "pages": pages,
            "full_text": "\n\n".join(full_text),
            "metadata": reader.metadata or {},
        }

    async def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse Word document."""
        doc = Document(str(file_path))

        paragraphs = []
        tables = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(
                    {"text": para.text, "style": para.style.name if para.style else None}
                )

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        full_text = "\n".join([p["text"] for p in paragraphs])

        return {
            "type": "docx",
            "filename": file_path.name,
            "paragraph_count": len(paragraphs),
            "paragraphs": paragraphs,
            "tables": tables,
            "full_text": full_text,
        }

    async def _parse_text(self, file_path: Path, doc_type: str = "text") -> Dict[str, Any]:
        """Parse text file with sampling for large logs."""
        file_size = file_path.stat().st_size
        max_size = 5 * 1024 * 1024  # 5MB

        if file_size > max_size and doc_type == "log":
            return await self._sample_large_log(file_path)

        async with aiofiles.open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = await f.read()

        lines = content.split("\n")

        return {
            "type": doc_type,
            "filename": file_path.name,
            "line_count": len(lines),
            "char_count": len(content),
            "full_text": content,
        }

    async def _sample_large_log(self, file_path: Path) -> Dict[str, Any]:
        """Sample the beginning and end of a large log file."""
        head_lines = []
        tail_lines = []
        num_lines = 5000

        async with aiofiles.open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            # Read head
            for _ in range(num_lines):
                line = await f.readline()
                if not line:
                    break
                head_lines.append(line)

            # Read tail
            file_size = file_path.stat().st_size
            # Seek back 512KB to find the last lines
            await f.seek(max(0, file_size - 1024 * 512))
            tail_content = await f.read()
            tail_lines = tail_content.splitlines()[-num_lines:]

        content = (
            "".join(head_lines)
            + "\n\n[... LOG SAMPLED DUE TO SIZE ...]\n\n"
            + "\n".join(tail_lines)
        )

        return {
            "type": "log_sampled",
            "filename": file_path.name,
            "is_sampled": True,
            "full_text": content,
            "message": f"Log file sampled (first and last {num_lines} lines).",
        }

    async def _parse_image(self, file_path: Path) -> Dict[str, Any]:
        """Parse image file (extract metadata)."""
        with Image.open(file_path) as img:
            return {
                "type": "image",
                "filename": file_path.name,
                "format": img.format,
                "size": img.size,
                "mode": img.mode,
                "file_path": str(file_path),
            }

    async def _parse_video(self, file_path: Path) -> Dict[str, Any]:
        """Extract keyframes from video for analysis."""
        cap = cv2.VideoCapture(str(file_path))
        fps = cap.get(cv2.CAP_PROP_FPS)
        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        duration = frame_count / fps if fps > 0 else 0

        # Extract a frame every 10 seconds, up to 10 frames
        interval = max(1, int(fps * 10))
        frames = []

        count = 0
        while cap.isOpened() and len(frames) < 10:
            ret, frame = cap.read()
            if not ret:
                break

            if count % interval == 0:
                # Save keyframe to a temporary directory
                temp_dir = Path(tempfile.gettempdir()) / "itsg33_keyframes"
                temp_dir.mkdir(exist_ok=True)

                frame_filename = f"{file_path.stem}_frame_{len(frames)}.jpg"
                frame_path = temp_dir / frame_filename
                cv2.imwrite(str(frame_path), frame)
                frames.append({"timestamp": count / fps, "path": str(frame_path)})

            count += 1

        cap.release()

        return {
            "type": "video",
            "filename": file_path.name,
            "duration_seconds": duration,
            "frame_count": frame_count,
            "keyframes": frames,
            "full_text": f"Video file: {file_path.name}. Duration: {duration:.2f}s. {len(frames)} keyframes extracted for analysis.",
        }

    def get_supported_extensions(self) -> List[str]:
        """Return list of supported file extensions."""
        return list(self.SUPPORTED_EXTENSIONS)

    async def extract_text(self, file_path: Path) -> str:
        """Extract plain text from a document."""
        parsed = await self.parse(file_path)
        if parsed and "full_text" in parsed:
            return parsed["full_text"]
        return ""
